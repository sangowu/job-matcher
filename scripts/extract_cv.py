#!/usr/bin/env python3
"""模块1：输入解析。

把 CV 文件（pdf/docx/txt/md）解析成纯文本，落盘并算稳定 hash，输出 JSON 元信息。
只负责"取文本"，不做任何字段理解（那是模块2）。

用法:
    python extract_cv.py <file_path>

输出(stdout, JSON):
    成功: {ok, source_type, char_count, cv_hash, text_path,
           cache_hit, cached_profile_path?, warnings}
    失败: {ok: false, error, source_type}

退出码: 0 成功 / 1 失败
"""
from __future__ import annotations

import hashlib
import json
import re
import sys
from pathlib import Path

# ── 路径约定：scripts/ 的上级是 skill 根，data/ 在根下 ──────────────────────────
SKILL_ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = SKILL_ROOT / "data"
CV_CACHE_DIR = DATA_DIR / "cv"
CV_TEXT_PATH = DATA_DIR / "cv_text.txt"

# ── 阈值 ────────────────────────────────────────────────────────────────────────
MAX_PAGES = 80           # PDF 超过此页数截断
MAX_CHARS = 200_000      # 文本超过此长度截断
SCAN_MIN_CHARS = 50      # 低于此长度疑似扫描件/空文档

SUPPORTED = {".pdf", ".docx", ".txt", ".md"}


def normalize_for_hash(text: str) -> str:
    """规范化文本用于算 hash —— 保证同一份 CV 多次解析 hash 稳定。

    压缩连续空白、统一换行、去首尾空白。注意：这只用于 hash，
    落盘的文本保留原始换行/bullet 结构供模块2 理解分节。
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_text_with_fallback(path: Path) -> str:
    """按编码兜底链读取纯文本文件。"""
    for enc in ("utf-8-sig", "utf-8", "gbk", "latin1"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    # latin1 理论上不会失败（单字节全覆盖），兜底再尝试 errors=replace
    return path.read_text(encoding="utf-8", errors="replace")


def extract_pdf(path: Path) -> tuple[str, list[str]]:
    """用 pdfplumber 逐页提取文本。"""
    import pdfplumber

    warnings: list[str] = []
    parts: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            pages = pdf.pages
            if len(pages) > MAX_PAGES:
                warnings.append(f"PDF 共 {len(pages)} 页，超过 {MAX_PAGES} 页，仅处理前 {MAX_PAGES} 页")
                pages = pages[:MAX_PAGES]
            for i, page in enumerate(pages, 1):
                txt = page.extract_text() or ""
                if not txt.strip():
                    warnings.append(f"第 {i} 页无文本")
                parts.append(txt)
    except Exception as e:  # 加密、损坏等
        msg = str(e).lower()
        if "password" in msg or "encrypt" in msg:
            raise RuntimeError("PDF 已加密，无法解析，请提供解密版本或粘贴文本") from e
        raise RuntimeError(f"PDF 解析失败：{e}") from e

    text = "\n".join(parts)
    if len(text.strip()) < SCAN_MIN_CHARS:
        warnings.append("提取文本极短，疑似扫描件（纯图片无文本层），建议改用可复制文本的版本或直接粘贴")
    return text, warnings


def extract_docx(path: Path) -> tuple[str, list[str]]:
    """用 python-docx 提取段落 + 表格单元格（很多模板把内容放表格里）。

    已知限制：python-docx 读不到文本框(textbox)里的文字。
    """
    import docx

    warnings: list[str] = []
    document = docx.Document(str(path))

    lines: list[str] = []
    for para in document.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    text = "\n".join(lines)
    if len(text.strip()) < SCAN_MIN_CHARS:
        warnings.append("DOCX 正文极短，可能内容在文本框中（python-docx 读不到），建议转 PDF 或直接粘贴文本")
    return text, warnings


def extract_plain(path: Path) -> tuple[str, list[str]]:
    """txt/md 直接读取（编码兜底链）。"""
    text = _read_text_with_fallback(path)
    warnings: list[str] = []
    if len(text.strip()) < SCAN_MIN_CHARS:
        warnings.append("文件内容极短")
    return text, warnings


def _fail(error: str, source_type: str = "unknown") -> None:
    print(json.dumps({"ok": False, "error": error, "source_type": source_type}))
    sys.exit(1)


def main() -> None:
    if len(sys.argv) != 2:
        _fail("用法: python extract_cv.py <file_path>")

    path = Path(sys.argv[1]).expanduser()
    if not path.exists():
        _fail(f"文件不存在: {path}")
    if not path.is_file():
        _fail(f"不是文件: {path}")

    suffix = path.suffix.lower()
    source_type = {".pdf": "pdf", ".docx": "docx", ".txt": "text", ".md": "text"}.get(suffix, "unknown")

    if suffix == ".doc":
        _fail("不支持旧版 .doc 格式，请另存为 .docx / PDF 或直接粘贴文本", "doc")
    if suffix not in SUPPORTED:
        _fail(f"不支持的格式 {suffix}，仅支持 pdf/docx/txt/md，或直接粘贴文本", source_type)

    try:
        if suffix == ".pdf":
            text, warnings = extract_pdf(path)
        elif suffix == ".docx":
            text, warnings = extract_docx(path)
        else:
            text, warnings = extract_plain(path)
    except RuntimeError as e:
        _fail(str(e), source_type)
        return  # 不会到这（_fail 已 exit），仅供类型检查

    # 长度上限截断
    if len(text) > MAX_CHARS:
        warnings.append(f"文本长度 {len(text)} 超过 {MAX_CHARS}，已截断")
        text = text[:MAX_CHARS]

    cv_hash = hashlib.sha256(normalize_for_hash(text).encode("utf-8")).hexdigest()[:16]

    # 落盘全文（保留原始结构），主上下文不接触全文
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    CV_TEXT_PATH.write_text(text, encoding="utf-8")

    # 缓存命中检查：同 CV 直接复用已抽取的 CVProfile
    cached_profile = CV_CACHE_DIR / f"{cv_hash}.json"
    cache_hit = cached_profile.exists()

    result = {
        "ok": True,
        "source_type": source_type,
        "char_count": len(text),
        "cv_hash": cv_hash,
        "text_path": str(CV_TEXT_PATH),
        "cache_hit": cache_hit,
        "warnings": warnings,
    }
    if cache_hit:
        result["cached_profile_path"] = str(cached_profile)

    # 统一纯 ASCII 输出（\u 转义），跨 console 编码安全，主 agent 解析 JSON 后还原中文
    print(json.dumps(result))


if __name__ == "__main__":
    main()
